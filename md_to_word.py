# -*- coding: utf-8 -*-
"""
Markdown转Word文档转换器
"""
import re
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def extract_tables(content):
    """提取并解析Markdown表格"""
    tables = []
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        if '|' in lines[i] and lines[i].strip().startswith('|'):
            # 找到表格开始
            table_lines = []
            header_line = i
            # 收集表格行
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i].strip())
                i += 1
            # 跳过分隔行
            if i < len(lines) and re.match(r'^[\|\s\-:]+$', lines[i]):
                i += 1
            tables.append((header_line, table_lines))
        else:
            i += 1
    return tables

def parse_table(table_lines):
    """解析表格为二维数组"""
    rows = []
    for line in table_lines:
        if re.match(r'^[\|\s\-:]+$', line):
            continue
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        rows.append(cells)
    return rows

def add_paragraph(doc, text, style='Normal', bold=False, font_size=None, color=None, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """添加段落"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold:
        run.bold = True
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    para.alignment = alignment
    return para

def add_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return heading

def add_table(doc, rows):
    """添加表格"""
    if not rows:
        return None
    # 计算列数
    max_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            # 设置字体
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    if i == 0:  # 表头加粗
                        run.bold = True
    return table

def md_to_word(md_content, doc):
    """将Markdown内容转换为Word文档"""
    # 处理代码块标记
    content = md_content
    
    # 处理表格（先提取，稍后添加）
    table_pattern = r'((?:\|[^\n]+\|\n)+)'
    
    # 按行处理
    lines = content.split('\n')
    i = 0
    pending_table = None
    
    while i < len(lines):
        line = lines[i]
        
        # 检测表格
        if '|' in line and line.strip().startswith('|'):
            # 收集整个表格
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            # 跳过分隔行
            if i < len(lines) and re.match(r'^[\|\s\-:]+$', lines[i]):
                i += 1
            
            # 解析并添加表格
            rows = parse_table(table_lines)
            if rows:
                add_table(doc, rows)
            doc.add_paragraph()  # 表格后添加空行
            continue
        
        # 处理标题
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            add_heading(doc, text, level)
            i += 1
            continue
        
        # 处理分隔线
        if re.match(r'^---+$', line.strip()) or re.match(r'^\*\*\*+$', line.strip()):
            i += 1
            continue
        
        # 处理引用块
        if line.strip().startswith('>'):
            quote_text = line.strip()[1:].strip()
            add_paragraph(doc, quote_text)
            i += 1
            continue
        
        # 处理无序列表
        list_match = re.match(r'^[-*]\s+(.+)$', line)
        if list_match:
            text = list_match.group(1)
            para = doc.add_paragraph(style='List Bullet')
            run = para.add_run(text)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            i += 1
            continue
        
        # 处理有序列表
        num_list_match = re.match(r'^\d+\.\s+(.+)$', line)
        if num_list_match:
            text = num_list_match.group(1)
            para = doc.add_paragraph(style='List Number')
            run = para.add_run(text)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            i += 1
            continue
        
        # 处理代码块
        if line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # 跳过结束标记
            # 添加代码块
            code_text = '\n'.join(code_lines)
            para = doc.add_paragraph()
            run = para.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            para.paragraph_format.left_indent = Inches(0.3)
            continue
        
        # 处理普通段落
        if line.strip():
            # 处理加粗和斜体
            text = line
            para = doc.add_paragraph()
            
            # 分割处理格式
            parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = para.add_run(part[1:-1])
                    run.italic = True
                elif part.startswith('`') and part.endswith('`'):
                    run = para.add_run(part[1:-1])
                    run.font.name = 'Consolas'
                else:
                    run = para.add_run(part)
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        else:
            # 空行，跳过
            pass
        
        i += 1

def convert_file(input_path, output_path):
    """转换单个文件"""
    print(f"正在读取: {input_path}")
    
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 创建Word文档
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 转换内容
    md_to_word(content, doc)
    
    # 保存
    doc.save(output_path)
    print(f"已保存: {output_path}")

if __name__ == '__main__':
    # 桌面路径
    desktop = os.path.join(os.path.expanduser("~"), "Desktop")
    
    # 转换项目文档
    input1 = r"C:\Users\1\WorkBuddy\20260407195418\美妆小程序项目文档.md"
    output1 = os.path.join(desktop, "i妆_美妆小程序项目文档.docx")
    convert_file(input1, output1)
    
    # 转换功能分析
    input2 = r"C:\Users\1\WorkBuddy\20260407195418\APP功能借鉴分析.md"
    output2 = os.path.join(desktop, "i妆_APP功能借鉴分析.docx")
    convert_file(input2, output2)
    
    print("\n✅ 转换完成！")
    print(f"📁 文件已保存到桌面:")
    print(f"   1. i妆_美妆小程序项目文档.docx")
    print(f"   2. i妆_APP功能借鉴分析.docx")
