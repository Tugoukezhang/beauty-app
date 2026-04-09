# -*- coding: utf-8 -*-
import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def read_markdown(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_table(doc, lines, start_idx):
    """解析Markdown表格"""
    # 找到表头和分隔符
    table_lines = []
    i = start_idx
    while i < len(lines):
        line = lines[i].strip()
        if '|' in line and not line.replace('|', '').replace('-', '').replace(':', '').strip():
            # 分隔符行，跳过
            i += 1
            continue
        if '|' in line:
            table_lines.append(line)
            i += 1
        else:
            break

    if len(table_lines) < 2:
        return i

    # 创建表格
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows.append(cells)

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = cell_text
            # 表头加粗
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    return start_idx + len(table_lines)

def md_to_doc(md_content, output_path):
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)

    lines = md_content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        # 跳过空行
        if not line:
            i += 1
            continue

        # 标题
        if line.startswith('# '):
            add_heading(doc, line[2:], 0)
            i += 1
        elif line.startswith('## '):
            add_heading(doc, line[3:], 1)
            i += 1
        elif line.startswith('### '):
            add_heading(doc, line[4:], 2)
            i += 1
        elif line.startswith('#### '):
            add_heading(doc, line[5:], 3)
            i += 1
        # 表格
        elif '|' in line:
            i = add_table(doc, lines, i)
        # 无序列表
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            i += 1
        # 有序列表
        elif re.match(r'^\d+\. ', line):
            p = doc.add_paragraph(line, style='List Number')
            i += 1
        # 引用
        elif line.startswith('> '):
            p = doc.add_paragraph(line[2:])
            p.italic = True
            i += 1
        # 水平线
        elif line.startswith('---'):
            i += 1
        # 普通段落
        else:
            # 处理加粗和斜体
            text = line
            p = doc.add_paragraph()
            # 简单处理**加粗**
            parts = re.split(r'\*\*(.*?)\*\*', text)
            for j, part in enumerate(parts):
                if j % 2 == 1:  # 加粗部分
                    run = p.add_run(part)
                    run.bold = True
                else:
                    run = p.add_run(part)
            i += 1

    doc.save(output_path)
    print(f'已保存: {output_path}')

# 桌面路径
desktop = r'C:\Users\1\Desktop'

# 源文件目录
source_dir = r'C:\Users\1\WorkBuddy\20260407195418\beauty-app-master'

# 要转换的文件
files = [
    ('小红书新娘妆30天选题计划.md', '小红书新娘妆30天选题计划.docx'),
    ('小红书美妆教程30天选题计划.md', '小红书美妆教程30天选题计划.docx'),
    ('小红书化妆博主起号手册.md', '小红书化妆博主起号手册.docx'),
]

for md_file, docx_file in files:
    md_path = os.path.join(source_dir, md_file)
    docx_path = os.path.join(desktop, docx_file)

    if os.path.exists(md_path):
        print(f'正在转换: {md_file}')
        content = read_markdown(md_path)
        md_to_doc(content, docx_path)
    else:
        print(f'文件不存在: {md_path}')
