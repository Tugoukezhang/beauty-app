#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown转Word文档（清理版）- 去除所有星号格式
"""
import io
import sys

# 设置stdout编码
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re
import os

def clean_text(text):
    """清理文本中的Markdown格式符号"""
    if not text:
        return ""
    
    # 移除 Markdown 强调符号（星号）
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # **bold** -> bold
    text = re.sub(r'\*(.+?)\*', r'\1', text)      # *italic* -> italic
    text = re.sub(r'~~(.+?)~~', r'\1', text)       # ~~strikethrough~~ -> text
    text = re.sub(r'`(.+?)`', r'\1', text)         # `code` -> code
    
    # 移除行首的列表标记
    text = re.sub(r'^[\-\*\+] ', '', text, flags=re.MULTILINE)  # - item -> item
    text = re.sub(r'^\d+\. ', '', text, flags=re.MULTILINE)     # 1. item -> item
    
    # 清理多余空格
    text = re.sub(r'\n{3,}', '\n\n', text)  # 多个换行 -> 两个换行
    
    return text.strip()

def is_table_row(line):
    """判断是否是一行表格分隔符"""
    return re.match(r'^[\|\-\s:]+$', line) is not None

def parse_markdown_table(lines, start_idx):
    """解析Markdown表格"""
    table_data = []
    
    # 跳过分隔符行
    idx = start_idx
    while idx < len(lines) and is_table_row(lines[idx]):
        idx += 1
    
    # 解析表头
    if idx < len(lines):
        headers = [cell.strip() for cell in lines[idx].split('|') if cell.strip()]
        table_data.append(headers)
        idx += 1
    
    # 跳过表头和内容之间的分隔符
    while idx < len(lines) and is_table_row(lines[idx]):
        idx += 1
    
    # 解析数据行
    while idx < len(lines):
        line = lines[idx].strip()
        if not line or is_table_row(line):
            idx += 1
            continue
        # 检查是否是非表格内容（标题或分隔符）
        if re.match(r'^#{1,6}\s', line):
            break
        if line.startswith('---'):
            break
            
        row = [cell.strip() for cell in line.split('|') if cell.strip()]
        if row:
            table_data.append(row)
        idx += 1
    
    return table_data, idx

def add_table_to_doc(doc, table_data):
    """将表格数据添加到文档"""
    if not table_data:
        return
    
    # 计算列数
    max_cols = max(len(row) for row in table_data)
    
    # 创建表格
    table = doc.add_table(rows=len(table_data), cols=max_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, row_data in enumerate(table_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                cell.text = clean_text(cell_text)
                
                # 表头加粗
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(10)

def add_code_block(doc, code_text):
    """添加代码块"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    paragraph.paragraph_format.left_indent = Inches(0.3)

def markdown_to_docx(md_content, output_path, title="文档"):
    """将Markdown内容转换为Word文档"""
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    
    lines = md_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # 跳过空行
        if not line:
            i += 1
            continue
        
        # 标题
        if line.startswith('# '):
            heading = doc.add_heading(level=0)
            heading.clear()
            run = heading.add_run(clean_text(line[2:]))
            run.font.size = Pt(22)
            run.font.bold = True
            i += 1
            
        elif line.startswith('## '):
            heading = doc.add_heading(level=1)
            heading.clear()
            run = heading.add_run(clean_text(line[3:]))
            run.font.size = Pt(18)
            i += 1
            
        elif line.startswith('### '):
            heading = doc.add_heading(level=2)
            heading.clear()
            run = heading.add_run(clean_text(line[4:]))
            run.font.size = Pt(14)
            i += 1
            
        elif line.startswith('#### '):
            heading = doc.add_heading(level=3)
            heading.clear()
            run = heading.add_run(clean_text(line[5:]))
            run.font.size = Pt(12)
            i += 1
        
        # 引用块
        elif line.startswith('>'):
            quote_text = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                text = lines[i].strip()[1:].strip()
                if text:
                    quote_text.append(text)
                i += 1
            if quote_text:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                run = p.add_run(' | '.join(quote_text))
                run.font.italic = True
                run.font.color.rgb = RGBColor(100, 100, 100)
        
        # 表格
        elif line.startswith('|'):
            table_data, i = parse_markdown_table(lines, i)
            add_table_to_doc(doc, table_data)
        
        # 分隔线
        elif line.startswith('---') or line.startswith('***') or line.startswith('___'):
            doc.add_paragraph('─' * 50)
            i += 1
        
        # 代码块
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # 跳过结束 ```
            if code_lines:
                add_code_block(doc, '\n'.join(code_lines))
        
        # 列表项（无序）
        elif re.match(r'^[\-\*] ', line):
            items = []
            while i < len(lines) and re.match(r'^[\-\*] ', lines[i].strip()):
                text = lines[i].strip()[2:].strip()
                items.append(clean_text(text))
                i += 1
            for item in items:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(item)
        
        # 列表项（有序）
        elif re.match(r'^\d+\. ', line):
            items = []
            while i < len(lines) and re.match(r'^\d+\. ', lines[i].strip()):
                text = re.sub(r'^\d+\. ', '', lines[i].strip())
                items.append(clean_text(text))
                i += 1
            for item in items:
                p = doc.add_paragraph(style='List Number')
                p.add_run(item)
        
        # 普通段落
        else:
            cleaned = clean_text(line)
            if cleaned:
                doc.add_paragraph(cleaned)
            i += 1
    
    # 保存文档
    doc.save(output_path)
    print(f"✓ 已保存: {output_path}")

def main():
    desktop = os.path.join(os.path.expanduser("~"), "Desktop")
    
    # 文件路径
    md_file1 = r"C:\Users\1\WorkBuddy\20260407195418\美妆小程序项目文档.md"
    md_file2 = r"C:\Users\1\WorkBuddy\20260407195418\APP功能借鉴分析.md"
    
    output1 = os.path.join(desktop, "i妆_美妆小程序项目文档_干净版.docx")
    output2 = os.path.join(desktop, "i妆_APP功能借鉴分析_干净版.docx")
    
    print("=" * 60)
    print("Markdown转Word文档（清理版）")
    print("=" * 60)
    
    # 转换文档1
    print(f"\n正在转换: {os.path.basename(md_file1)}")
    with open(md_file1, 'r', encoding='utf-8') as f:
        content1 = f.read()
    markdown_to_docx(content1, output1, "i妆_美妆小程序项目文档")
    
    # 转换文档2
    print(f"\n正在转换: {os.path.basename(md_file2)}")
    with open(md_file2, 'r', encoding='utf-8') as f:
        content2 = f.read()
    markdown_to_docx(content2, output2, "i妆_APP功能借鉴分析")
    
    print("\n" + "=" * 60)
    print("转换完成！请查看桌面上的干净版Word文档。")
    print("=" * 60)

if __name__ == "__main__":
    main()
