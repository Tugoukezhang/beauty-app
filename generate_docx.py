#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Markdown转Word文档工具"""

import re
import sys
from pathlib import Path

try:
    import markdown
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("请先安装依赖：pip install python-docx markdown")
    sys.exit(1)


def md_to_docx(md_file_path, docx_file_path):
    """将Markdown文件转换为Word文档"""

    # 读取Markdown文件
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # 创建Word文档
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)

    # 处理Markdown内容
    lines = md_content.split('\n')
    i = 0
    in_table = False
    table_data = []

    while i < len(lines):
        line = lines[i]

        # 跳过HTML注释和模板标记
        if line.strip().startswith('<') and line.strip().endswith('>'):
            i += 1
            continue
        if line.strip().startswith('---') and line.strip().endswith('---'):
            i += 1
            continue

        # 标题处理
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=3)

        # 表格处理
        elif line.startswith('|') and line.strip().endswith('|'):
            # 检测表头分隔行
            if i + 1 < len(lines) and re.match(r'\|[\s\-:|]+\|', lines[i + 1]):
                # 这是一个表格
                table_lines = [line]
                i += 1
                # 跳过分隔行
                i += 1
                # 收集表格行
                while i < len(lines) and lines[i].startswith('|') and lines[i].strip().endswith('|'):
                    table_lines.append(lines[i])
                    i += 1
                # 创建表格
                if table_lines:
                    # 解析表格
                    rows_data = []
                    for table_line in table_lines:
                        cells = [cell.strip() for cell in table_line.split('|')[1:-1]]
                        if cells and any(cells):
                            rows_data.append(cells)

                    if rows_data:
                        # 创建表格（行数+表头）
                        table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
                        table.style = 'Table Grid'

                        for row_idx, row_data in enumerate(rows_data):
                            for col_idx, cell_text in enumerate(row_data):
                                cell = table.rows[row_idx].cells[col_idx]
                                cell.text = cell_text
                                # 表头加粗
                                if row_idx == 0:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.bold = True

                continue  # 继续处理下一行

        # 代码块处理
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                code_block = doc.add_paragraph()
                code_run = code_block.add_run('\n'.join(code_lines))
                code_run.font.name = 'Consolas'
                code_run.font.size = Pt(9)
                code_block.paragraph_format.left_indent = Inches(0.5)
        # 分隔线
        elif line.strip() == '---':
            doc.add_paragraph('─' * 50)
        # 空行
        elif line.strip() == '':
            doc.add_paragraph()
        # 普通文本处理
        else:
            # 处理Markdown格式
            text = line

            # 处理加粗 **text**
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)

            # 处理行内代码 `code`
            text = re.sub(r'`(.+?)`', r'\1', text)

            # 处理链接 [text](url)
            text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)

            # 添加段落
            if text.strip():
                p = doc.add_paragraph(text)
                p.paragraph_format.left_indent = Inches(0)

        i += 1

    # 保存文档
    doc.save(docx_file_path)
    print(f"已生成: {docx_file_path}")


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("用法: python md_to_docx.py <输入.md> <输出.docx>")
        sys.exit(1)

    md_file = sys.argv[1]
    docx_file = sys.argv[2]

    md_to_docx(md_file, docx_file)
